import io
import logging

logger = logging.getLogger("ai-service.parser")

MAX_CHARS = 20000


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Pull plain text out of a resume. Returns "" when the file cannot be read
    (scanned image PDFs, corrupt files) so the caller can fail gracefully
    instead of sending garbage to the model.
    """
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        text = _from_pdf(file_bytes)
    elif name.endswith(".docx"):
        text = _from_docx(file_bytes)
    elif name.endswith(".doc"):
        # Legacy binary .doc has no pure-python reader worth shipping.
        logger.warning("Legacy .doc uploaded; cannot extract text")
        return ""
    else:
        logger.warning("Unsupported resume type: %s", filename)
        return ""

    return _clean(text)


def _from_pdf(data: bytes) -> str:
    try:
        import pdfplumber

        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            # Ten pages is far more than any real resume; the cap stops a
            # 200-page upload from blowing up latency and token cost.
            for page in pdf.pages[:10]:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except Exception:
        logger.exception("PDF extraction failed")
        return ""


def _from_docx(data: bytes) -> str:
    try:
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]

        # Plenty of resumes lay everything out inside tables.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts)
    except Exception:
        logger.exception("DOCX extraction failed")
        return ""


def _clean(text: str) -> str:
    if not text:
        return ""

    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    cleaned = "\n".join(lines)

    if len(cleaned) > MAX_CHARS:
        cleaned = cleaned[:MAX_CHARS]

    return cleaned